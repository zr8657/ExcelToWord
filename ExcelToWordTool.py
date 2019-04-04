# -*- coding: utf-8 -*-
from imp import reload

from openpyxl import load_workbook
from docx import Document
import time
import sys
import easygui as g
from PyQt5 import QtWidgets
from PyQt5.QtWidgets import QFileDialog, QWidget
import PyQt5.QtCore
import re


class MyWindow(QWidget):
    def __init__(self):
        super(MyWindow, self).__init__()
        file_path, filetype = QFileDialog.getOpenFileName(self, "选择文件", "/", "All Files (*);;Text Files (*.txt)")
        print(file_path)  # 打印文件全部路径（包括文件名和后缀名）
        # 获取文件名
        file_name = re.findall(r'[^\\/:*?"<>|\r\n]+$', file_path)
        file_name = re.findall(r'(.+?)\.xlsx', file_name[0])
        print(file_name[0] + '.xlsx')
        try:
            reload(sys)
            # 开始时间
            startTime = time.time()

            # 读取excel xlsx文件
            wb = load_workbook(file_path)

            # 获取所有sheet页名字
            xl_sheet_names = wb.get_sheet_names()

            # 定位到相应sheet页,[0]为sheet页索引
            xl_sheet = wb.get_sheet_by_name(xl_sheet_names[0])

            # 获取行列数
            excel_row = xl_sheet.max_row
            excel_column = xl_sheet.max_column

            # word文档
            document = Document()

            document.add_heading('危险化学品表格', 0)

            table = document.add_table(rows=1, cols=excel_column, style="Table Grid")
            # hdr_cells = table.rows[0].cells
            # for num in range(0, excel_column):
            #     hdr_cells[num].text = u'' + str(xl_sheet.rows[num].value)

            # 取excel第一页第一张表
            i = 0

            # 写入word

            # 将excel表格装入itercars操作
            itercars = iter(xl_sheet.rows)
            # 列名
            hdr_cells = table.rows[0].cells
            for row in xl_sheet.rows:
                for num in range(0, excel_column):
                    hdr_cells[num].text = u'' + str(row[num].value)
                break

            # 进入第二行开始循环插入
            next(itercars)
            for row in itercars:
                row_cells = table.add_row().cells
                for num in range(0, excel_column):
                    content = str(row[num].value)
                    content = "" if content == "None" else content
                    row_cells[num].text = u'' + content

            document.add_page_break()
            # document.save('demo.docx')
            docx_path = file_path.replace('xlsx', 'docx')
            document.save(docx_path)
            title = g.msgbox(msg="                                     成功！", title="Success", ok_button="确定")
        except Exception as e:
            print(e)
            title = g.msgbox(msg="               生成失败:" + e, title="Error", ok_button="确定")


if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    myshow = MyWindow()
    myshow.show()
    myshow.close()
    sys.exit()
